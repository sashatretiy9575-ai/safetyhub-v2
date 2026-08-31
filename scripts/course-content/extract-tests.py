from __future__ import annotations

import argparse
import hashlib
import json
import re
import uuid
from pathlib import Path

from docx import Document


NAMESPACE = uuid.UUID("8a87d7f2-9345-5bd4-9dc9-eefdf3b9313e")
DOCX_SHA256 = "3b36bba9f031c233cd10a4a80e1f7dd2a89af523b5df2145056dc7366086ce00"
UPDATED_AT = "2026-08-25T00:00:00.000Z"
EFFECTIVE_DATE = "2026-08-21"
LEGACY_WORDING = "Можно ли определять шаг анкеров по универсальной цифре из старой презентации?"
REPLACEMENT_WORDING = "Можно ли определять шаг анкеров по одной универсальной цифре?"


COURSES = [
    {
        "heading": "ПЛОТНИК",
        "slug": "plotnik",
        "title": "Плотник",
        "description": "Безопасные методы работы плотника: организация рабочего места, ручной и электроинструмент, СИЗ и действия при неисправностях.",
        "icon": "hammer",
        "displayOrder": 1,
        "sourceFilename": "01_Плотник.pptx",
        "sourceSha256": "d3b84cbe64a5376c61168ee798b5130c4e593ba39c302f785a13293125766948",
        "pageCount": 25,
    },
    {
        "heading": "АРМАТУРЩИК",
        "slug": "armaturshchik",
        "title": "Арматурщик",
        "description": "Безопасная подготовка, обработка, перемещение и монтаж арматуры с контролем инструмента, каркаса и рабочей зоны.",
        "icon": "hard-hat",
        "displayOrder": 2,
        "sourceFilename": "02_Арматурщик.pptx",
        "sourceSha256": "480cd55c460114d9210fae0a1b7232a71345587c676cdecc644351ec2497ffbe",
        "pageCount": 31,
    },
    {
        "heading": "ЛЕСОМОНТАЖНЫЕ РАБОТЫ",
        "slug": "lesomontazhnye-raboty",
        "title": "Лесомонтажные работы",
        "description": "Безопасный монтаж и эксплуатация лесов: проект, основание, связи, настил, ограждения, доступ, допуск и контроль условий.",
        "icon": "ladder",
        "displayOrder": 3,
        "sourceFilename": "03_Лесомонтажные_работы.pptx",
        "sourceSha256": "7b8b4b247ba1fe3ff3480ded219417ece9710daa0fd12875f81db10d058228d6",
        "pageCount": 42,
    },
    {
        "heading": "БИОТ",
        "slug": "biot",
        "title": "БИОТ",
        "description": "Базовые требования безопасности и охраны труда: опасности, оценка риска, обучение, допуск, средства защиты и аварийные действия.",
        "icon": "shield",
        "displayOrder": 4,
        "sourceFilename": "04_БИОТ.pptx",
        "sourceSha256": "aaca0cb2b612774e7d574f66d0bf2c103536333bb575dca43db4cf657b601c8f",
        "pageCount": 59,
    },
    {
        "heading": "ПОЖАРНАЯ БЕЗОПАСНОСТЬ",
        "slug": "pozharnaya-bezopasnost",
        "title": "Пожарная безопасность",
        "description": "Профилактика пожаров, оповещение, эвакуация, применение первичных средств тушения и безопасные действия при пожаре.",
        "icon": "fire",
        "displayOrder": 5,
        "sourceFilename": "05_Пожарная_безопасность.pptx",
        "sourceSha256": "eb0c64b3aa83a3d63d74631d899def9f7683d06858baca516cc8fc254b46d68a",
        "pageCount": 41,
    },
]

EXPECTED_KEYS = {
    "plotnik": [
        list("ВАГБВГАВБГ"),
        list("БГАВБГВАГБ"),
        list("АВБГВАГБАВ"),
    ],
    "armaturshchik": [
        list("ГБВАГВБГАВ"),
        list("АБГВБАГБВГ"),
        list("ВАБВГАВБГА"),
    ],
    "lesomontazhnye-raboty": [
        list("АГВБГАВГБВ"),
        list("ВБГАБГВБАГ"),
        list("БВАГВАБВГА"),
    ],
    "biot": [
        list("БВГАВГБАГВ"),
        list("ГАБГВБАГВБ"),
        list("АГВБАВГАБВ"),
    ],
    "pozharnaya-bezopasnost": [
        list("ГВАБВГАВГБ"),
        list("БАГБВГБАГВ"),
        list("ВБАВГБАВГА"),
    ],
}

LETTER_TO_INDEX = {"А": 0, "Б": 1, "В": 2, "Г": 3}


def normalize(value: str) -> str:
    return re.sub(r"[\s\u00a0]+", " ", value).strip()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def stable_uuid(name: str) -> str:
    return str(uuid.uuid5(NAMESPACE, name))


def canonical_hash(value: object) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def database_content_projection(course: dict[str, object]) -> dict[str, object]:
    """Mirror private.course_content_hash_v3's cross-runtime canonical projection."""
    policy = course["policy"]
    presentation = course["presentation"]
    return {
        "slug": course["slug"],
        "title": course["title"],
        "description": course["description"],
        "icon": course["icon"],
        "displayOrder": course["displayOrder"],
        "presentationSha256": presentation["sha256"],
        "presentationPageCount": presentation["pageCount"],
        "durationMinutes": policy["durationMinutes"],
        "passScore": policy["passScore"],
        "attemptsPerCalendarDay": policy["attemptsPerCalendarDay"],
        "attemptResetTimezone": policy["resetTimezone"],
        "questionVariants": course["variants"],
        "seo": course["seo"],
        "jurisdiction": course["jurisdiction"],
        "effectiveDate": course["effectiveDate"],
        "sources": course["sources"],
    }


def write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def extract_sections(document: Document) -> dict[str, list[list[str]]]:
    headings = {course["heading"]: course["slug"] for course in COURSES}
    sections: dict[str, list[list[str]]] = {course["slug"]: [] for course in COURSES}
    current_slug: str | None = None
    current_variant: int | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_variant, current_lines
        if current_slug is None or current_variant is None:
            return
        if len(current_lines) != 50:
            raise ValueError(
                f"{current_slug} variant {current_variant}: expected 50 content lines, got {len(current_lines)}"
            )
        expected_next = len(sections[current_slug]) + 1
        if current_variant != expected_next:
            raise ValueError(f"{current_slug}: expected variant {expected_next}, got {current_variant}")
        sections[current_slug].append(current_lines)
        current_variant = None
        current_lines = []

    for paragraph in document.paragraphs:
        text = normalize(paragraph.text)
        if not text:
            continue
        upper = text.upper()
        if upper in {"ОТВЕТЫ", "СЛУЖЕБНЫЙ РАЗДЕЛ", "КЛЮЧИ ПРАВИЛЬНЫХ ОТВЕТОВ"}:
            flush()
            break
        if upper in headings:
            flush()
            current_slug = headings[upper]
            continue
        variant_match = re.fullmatch(r"Вариант\s+([123])", text, flags=re.IGNORECASE)
        if variant_match:
            flush()
            if current_slug is None:
                raise ValueError(f"Variant heading before course heading: {text}")
            current_variant = int(variant_match.group(1))
            continue
        if current_variant is not None:
            if text == "Выберите один правильный ответ в каждом вопросе.":
                continue
            current_lines.append(text)

    for slug, variants in sections.items():
        if len(variants) != 3:
            raise ValueError(f"{slug}: expected three variants, got {len(variants)}")
    return sections


def extract_answer_keys(document: Document) -> dict[str, list[list[str]]]:
    key_tables = document.tables[-5:]
    if len(key_tables) != 5:
        raise ValueError("Expected five answer-key tables")
    result: dict[str, list[list[str]]] = {}
    for course, table in zip(COURSES, key_tables, strict=True):
        if len(table.rows) != 4 or len(table.columns) != 11:
            raise ValueError(f"{course['slug']}: malformed answer-key table")
        variants: list[list[str]] = []
        for row_index in range(1, 4):
            variant_number = normalize(table.cell(row_index, 0).text)
            if variant_number != str(row_index):
                raise ValueError(f"{course['slug']}: expected key row {row_index}, got {variant_number!r}")
            letters = [normalize(table.cell(row_index, column).text).upper() for column in range(1, 11)]
            if any(letter not in LETTER_TO_INDEX for letter in letters):
                raise ValueError(f"{course['slug']} variant {row_index}: unsupported answer letter")
            variants.append(letters)
        result[course["slug"]] = variants
    if result != EXPECTED_KEYS:
        raise ValueError("Answer-key tables do not match the approved control matrix")
    return result


def build_course(
    course: dict[str, object],
    section_variants: list[list[str]],
    answer_keys: list[list[str]],
    snapshot_root: Path,
) -> dict[str, object]:
    slug = str(course["slug"])
    course_id = stable_uuid(f"course:{slug}")
    presentation_id = stable_uuid(f"presentation:{slug}:v1")
    course_dir = snapshot_root / slug
    presentation_manifest_path = course_dir / "presentation-manifest.json"
    if not presentation_manifest_path.exists():
        raise FileNotFoundError(f"Missing presentation manifest: {presentation_manifest_path}")
    presentation_manifest = json.loads(presentation_manifest_path.read_text(encoding="utf-8"))
    if presentation_manifest["pageCount"] != course["pageCount"]:
        raise ValueError(f"{slug}: presentation page-count mismatch")

    variants: list[dict[str, object]] = []
    correction_count = 0
    for variant_index, (lines, key_letters) in enumerate(zip(section_variants, answer_keys, strict=True), start=1):
        variant_id = stable_uuid(f"variant:{slug}:{variant_index}")
        questions: list[dict[str, object]] = []
        for question_offset in range(10):
            question_number = question_offset + 1
            base = question_offset * 5
            question_text = lines[base]
            if question_text == LEGACY_WORDING:
                if slug != "lesomontazhnye-raboty" or variant_index != 2 or question_number != 5:
                    raise ValueError("Legacy wording appeared outside the approved correction location")
                question_text = REPLACEMENT_WORDING
                correction_count += 1
            option_texts = lines[base + 1 : base + 5]
            question_id = stable_uuid(f"question:{slug}:{variant_index}:{question_number}")
            options = [
                {
                    "id": stable_uuid(f"option:{slug}:{variant_index}:{question_number}:{option_index}"),
                    "text": option_text,
                    "displayOrder": option_index,
                }
                for option_index, option_text in enumerate(option_texts, start=1)
            ]
            correct_index = LETTER_TO_INDEX[key_letters[question_offset]]
            questions.append(
                {
                    "id": question_id,
                    "text": question_text,
                    "displayOrder": question_number,
                    "options": options,
                    "correctOptionId": options[correct_index]["id"],
                    "explanation": "",
                }
            )
        variants.append(
            {
                "id": variant_id,
                "variantNumber": variant_index,
                "questions": questions,
            }
        )

    if slug == "lesomontazhnye-raboty":
        if correction_count != 1:
            raise ValueError(f"Expected exactly one wording correction, got {correction_count}")
    elif correction_count != 0:
        raise ValueError(f"{slug}: unexpected wording correction")

    pdf_sha = presentation_manifest["sha256"]
    thumbnail_sha = presentation_manifest["thumbnailSha256"]
    presentation = {
        "id": presentation_id,
        "sourceFilename": course["sourceFilename"],
        "operatorPptx": f"content/source-materials/derived/{slug}/presentation.pptx",
        "file": "presentation.pdf",
        "thumbnail": "thumbnail.webp",
        "storageBucket": "course-presentations",
        "storagePath": f"{course_id}/{presentation_id}/{pdf_sha}.pdf",
        "thumbnailPath": f"{course_id}/{presentation_id}/{pdf_sha}-thumb.webp",
        "mimeType": "application/pdf",
        "byteSize": presentation_manifest["byteSize"],
        "pageCount": presentation_manifest["pageCount"],
        "aspectRatio": "16:9",
        "sha256": pdf_sha,
        "thumbnailSha256": thumbnail_sha,
        "notesIncluded": False,
    }

    result: dict[str, object] = {
        "schemaVersion": 1,
        "id": course_id,
        "slug": slug,
        "title": course["title"],
        "description": course["description"],
        "icon": course["icon"],
        "displayOrder": course["displayOrder"],
        "updatedAt": UPDATED_AT,
        "jurisdiction": "Республика Казахстан",
        "effectiveDate": EFFECTIVE_DATE,
        "seo": {},
        "sources": [],
        "policy": {
            "durationMinutes": 15,
            "passScore": 7,
            "questionCount": 10,
            "variantCount": 3,
            "attemptsPerCalendarDay": 8,
            "resetTimezone": "Asia/Oral",
        },
        "presentation": presentation,
        "variants": variants,
        "sourceMaterials": {
            "presentation": {
                "filename": course["sourceFilename"],
                "sha256": course["sourceSha256"],
            },
            "tests": {
                "filename": "06_Тестовые_задания.docx",
                "sha256": DOCX_SHA256,
            },
        },
    }
    result["dbContentHash"] = canonical_hash(database_content_projection(result))
    result["snapshotContentHash"] = canonical_hash(result)
    # Compatibility alias: contentHash always means the authoritative DB hash.
    result["contentHash"] = result["dbContentHash"]
    return result


def validate_catalog(courses: list[dict[str, object]]) -> dict[str, int]:
    slugs = [str(course["slug"]) for course in courses]
    if slugs != [str(course["slug"]) for course in COURSES]:
        raise ValueError("Course order or slug set is not canonical")
    all_ids: set[str] = set()
    variant_count = question_count = option_count = correct_count = page_count = 0
    for course in courses:
        page_count += int(course["presentation"]["pageCount"])
        variants = course["variants"]
        if len(variants) != 3:
            raise ValueError(f"{course['slug']}: expected three variants")
        variant_count += len(variants)
        for variant in variants:
            questions = variant["questions"]
            if len(questions) != 10:
                raise ValueError(f"{course['slug']} variant {variant['variantNumber']}: expected ten questions")
            question_count += len(questions)
            for question in questions:
                if not question["text"].strip() or re.search(r"<\s*script", question["text"], flags=re.IGNORECASE):
                    raise ValueError(f"{course['slug']}: unsafe or empty question")
                options = question["options"]
                if len(options) != 4 or any(not option["text"].strip() for option in options):
                    raise ValueError(f"{course['slug']}: expected four non-empty options")
                option_ids = [option["id"] for option in options]
                if question["correctOptionId"] not in option_ids:
                    raise ValueError(f"{course['slug']}: correct option does not belong to question")
                option_count += len(options)
                correct_count += 1
                for identifier in [question["id"], *option_ids]:
                    if identifier in all_ids:
                        raise ValueError(f"Duplicate stable identifier: {identifier}")
                    all_ids.add(identifier)
        for identifier in [course["id"], course["presentation"]["id"], *[variant["id"] for variant in variants]]:
            if identifier in all_ids:
                raise ValueError(f"Duplicate stable identifier: {identifier}")
            all_ids.add(identifier)

    totals = {
        "courseCount": len(courses),
        "presentationCount": len(courses),
        "presentationPageCount": page_count,
        "variantCount": variant_count,
        "questionCount": question_count,
        "optionCount": option_count,
        "correctAnswerCount": correct_count,
    }
    expected_totals = {
        "courseCount": 5,
        "presentationCount": 5,
        "presentationPageCount": 198,
        "variantCount": 15,
        "questionCount": 150,
        "optionCount": 600,
        "correctAnswerCount": 150,
    }
    if totals != expected_totals:
        raise ValueError(f"Unexpected catalog totals: {totals}")
    return totals


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract the approved SafetyHub test bank into deterministic snapshots.")
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--snapshot-root", type=Path, default=Path("content/snapshots/courses"))
    args = parser.parse_args()

    docx_path = args.docx.resolve()
    source_dir = args.source_dir.resolve()
    snapshot_root = args.snapshot_root.resolve()
    if sha256(docx_path) != DOCX_SHA256:
        raise ValueError("Test DOCX SHA-256 does not match the approved source")
    for course in COURSES:
        source_path = source_dir / str(course["sourceFilename"])
        actual = sha256(source_path)
        if actual != course["sourceSha256"]:
            raise ValueError(f"{source_path.name}: source SHA-256 changed ({actual})")

    document = Document(docx_path)
    sections = extract_sections(document)
    answer_keys = extract_answer_keys(document)
    courses = [
        build_course(course, sections[str(course["slug"])], answer_keys[str(course["slug"])], snapshot_root)
        for course in COURSES
    ]
    totals = validate_catalog(courses)
    for course in courses:
        write_json(snapshot_root / str(course["slug"]) / "course.json", course)

    catalog_without_hash = {
        "schemaVersion": 1,
        "catalogVersion": "2026-08-25-new-five-course-catalog",
        "updatedAt": UPDATED_AT,
        "policy": {
            "durationMinutes": 15,
            "passScore": 7,
            "questionCount": 10,
            "variantCount": 3,
            "attemptsPerCalendarDay": 8,
            "resetTimezone": "Asia/Oral",
        },
        "totals": totals,
        "courses": [
            {
                "id": course["id"],
                "slug": course["slug"],
                "title": course["title"],
                "displayOrder": course["displayOrder"],
                "contentHash": course["contentHash"],
                "dbContentHash": course["dbContentHash"],
                "snapshotContentHash": course["snapshotContentHash"],
                "presentationSha256": course["presentation"]["sha256"],
                "pageCount": course["presentation"]["pageCount"],
            }
            for course in courses
        ],
        "sourceDocument": {
            "filename": docx_path.name,
            "sha256": DOCX_SHA256,
        },
        "answerKeyMatrix": EXPECTED_KEYS,
        "approvedWordingCorrection": {
            "courseSlug": "lesomontazhnye-raboty",
            "variantNumber": 2,
            "questionNumber": 5,
            "sourceTextSha256": hashlib.sha256(LEGACY_WORDING.encode("utf-8")).hexdigest(),
            "after": REPLACEMENT_WORDING,
            "legacyReferenceRemoved": True,
        },
    }
    catalog_without_hash["catalogChecksum"] = hashlib.sha256(
        ",".join(str(course["dbContentHash"]) for course in courses).encode("utf-8")
    ).hexdigest()
    catalog_without_hash["catalogHash"] = canonical_hash(catalog_without_hash)
    write_json(snapshot_root / "catalog.json", catalog_without_hash)
    print(json.dumps(totals, ensure_ascii=False, sort_keys=True))


if __name__ == "__main__":
    main()
